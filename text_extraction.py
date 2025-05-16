import io
import re
import PyPDF2
from docx import Document
import streamlit as st

# Add to text_extraction.py
def extract_text_from_pdf(pdf_file):
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages[:5]:  # Limit to first 5 pages
            try:
                text += page.extract_text() or ""  # Handle None returns
            except Exception as page_error:
                st.warning(f"Couldn't read a page: {str(page_error)}")
                continue
                
        if not text.strip():
            raise ValueError("PDF appears to be image-based or empty")
        return clean_text(text)
    except Exception as e:
        st.error(f"Critical PDF error: {str(e)}")
        return None

def extract_text_from_docx(docx_file):
    """
    Extract text from a DOCX file.
    
    Args:
        docx_file: A file object containing the DOCX
        
    Returns:
        str: Extracted text from the DOCX
    """
    try:
        # Create a temporary file-like object
        docx_bytes = io.BytesIO(docx_file.getvalue())
        
        # Load the document
        doc = Document(docx_bytes)
        
        # Extract text from each paragraph
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        # Join the text and clean it
        text = '\n'.join(full_text)
        text = clean_text(text)
        
        return text
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {str(e)}")
        return None

def clean_text(text):
    """
    Clean and normalize extracted text.
    
    Args:
        text: Raw text to clean
        
    Returns:
        str: Cleaned text
    """
    if not text:
        return ""
        
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove unusual characters and normalize
    text = re.sub(r'[^\x00-\x7F]+', '', text)
    
    # Trim leading/trailing whitespace
    text = text.strip()
    
    return text
